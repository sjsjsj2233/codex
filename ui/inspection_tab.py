"""
네트워크 장비 점검 보고서 생성 탭
show 명령어 텍스트 파일 → PDF / Word / Excel 보고서 자동 생성
"""
import os
from datetime import datetime

from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGroupBox,
    QLabel, QLineEdit, QPushButton, QListWidget,
    QListWidgetItem, QFileDialog, QMessageBox,
    QSplitter, QProgressBar, QScrollArea, QFrame,
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont

from core.inspection_parser import parse_file, DeviceInspection


# ── 상태 색상 ──────────────────────────────────────────────────────────────
_STATUS_COLOR = {'정상': '#16a34a', '주의': '#ea580c', '경고': '#dc2626', '파일 오류': '#6b7280'}

# ── 기본 점검 정보 항목 ────────────────────────────────────────────────────
_DEFAULT_FIELDS = [
    ('고객사',    '예: (주)ABC 네트워크', ''),
    ('사이트명',  '예: 본사 데이터센터', ''),
    ('고객 담당', '예: 홍길동', ''),
    ('점검일',    '',  datetime.now().strftime('%Y-%m-%d')),
    ('점검 회사', '예: NetOps 주식회사', ''),
    ('점검 목적', '', '정기 네트워크 장비 점검'),
    ('점검 범위', '', '네트워크 핵심 장비 (Router, Switch, Firewall)'),
]


# ── 백그라운드 생성 스레드 ─────────────────────────────────────────────────
class _ReportWorker(QThread):
    progress = pyqtSignal(str)
    done     = pyqtSignal(str)
    error    = pyqtSignal(str)

    def __init__(self, fmt, save_path, files, info):
        super().__init__()
        self.fmt       = fmt        # 'pdf', 'word', 'excel'
        self.save_path = save_path
        self.files     = files
        self.info      = info       # {'items': [{'label':..,'value':..}, ...]}

    def run(self):
        try:
            self.progress.emit('파일 파싱 중…')
            devices = []
            for fp in self.files:
                self.progress.emit(f'파싱: {os.path.basename(fp)}')
                devices.append(parse_file(fp))

            if self.fmt == 'pdf':
                self.progress.emit('PDF 생성 중…')
                _build_pdf(self.save_path, devices, self.info)
            elif self.fmt == 'word':
                self.progress.emit('Word 생성 중…')
                _build_word(self.save_path, devices, self.info)
            else:
                self.progress.emit('Excel 생성 중…')
                _build_excel(self.save_path, devices, self.info)

            self.done.emit(self.save_path)
        except Exception as ex:
            self.error.emit(str(ex))


# ═══════════════════════════════════════════════════════════════════════════
# 동적 정보 행 위젯
# ═══════════════════════════════════════════════════════════════════════════
class _InfoRow(QWidget):
    """라벨 + 값 입력 + 삭제 버튼 한 행"""
    delete_clicked = pyqtSignal(object)

    def __init__(self, label='', value='', placeholder='', parent=None):
        super().__init__(parent)
        h = QHBoxLayout(self)
        h.setContentsMargins(0, 2, 0, 2)
        h.setSpacing(6)

        self.label_edit = QLineEdit(label)
        self.label_edit.setPlaceholderText('항목명')
        self.label_edit.setFixedWidth(110)
        self.label_edit.setFixedHeight(26)

        self.value_edit = QLineEdit(value)
        self.value_edit.setPlaceholderText(placeholder or '값 입력')
        self.value_edit.setFixedHeight(26)

        del_btn = QPushButton('✕')
        del_btn.setFixedSize(24, 24)
        del_btn.setStyleSheet(
            'QPushButton{background:#ef4444;color:white;border-radius:4px;font-size:10px;}'
            'QPushButton:hover{background:#dc2626;}'
        )
        del_btn.clicked.connect(lambda: self.delete_clicked.emit(self))

        h.addWidget(self.label_edit)
        h.addWidget(self.value_edit, 1)
        h.addWidget(del_btn)

    def get_pair(self):
        return self.label_edit.text().strip(), self.value_edit.text().strip()


# ═══════════════════════════════════════════════════════════════════════════
# UI
# ═══════════════════════════════════════════════════════════════════════════
class InspectionTab:
    def __init__(self, parent=None):
        self._widget = _InspectionWidget(parent)

    def as_widget(self):
        return self._widget


class _InspectionWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._worker   = None
        self._info_rows = []
        self._build_ui()

    def _build_ui(self):
        from ui.report_tab import _Header
        self.setObjectName('inspectionWidget')
        self.setStyleSheet('#inspectionWidget { background: #f1f5f9; }')
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        root.addWidget(_Header(
            '점검 보고서',
            'show 명령어 파일 자동 분석 · PDF / Word / Excel 점검 보고서 생성',
            '#0f172a', '#ea580c',
        ))

        body = QWidget()
        body.setObjectName('inspectionBody')
        body.setStyleSheet('#inspectionBody { background: transparent; }')
        bv = QVBoxLayout(body)
        bv.setContentsMargins(10, 10, 10, 10)
        bv.setSpacing(8)
        root.addWidget(body, 1)

        splitter = QSplitter(Qt.Horizontal)

        # ── 왼쪽: 점검 정보 (동적) ──────────────────────────────
        left = QWidget()
        lv   = QVBoxLayout(left)
        lv.setContentsMargins(0, 0, 6, 0)
        lv.setSpacing(6)

        info_box = QGroupBox("📋 점검 정보")
        info_box_v = QVBoxLayout(info_box)
        info_box_v.setSpacing(4)

        # 열 헤더
        col_header = QHBoxLayout()
        col_header.setContentsMargins(0, 0, 0, 0)
        lbl_h1 = QLabel('항목명')
        lbl_h1.setFixedWidth(110)
        lbl_h1.setStyleSheet('color:#64748b;font-size:10px;font-weight:bold')
        lbl_h2 = QLabel('내용')
        lbl_h2.setStyleSheet('color:#64748b;font-size:10px;font-weight:bold')
        col_header.addWidget(lbl_h1)
        col_header.addWidget(lbl_h2)
        col_header.addSpacing(30)
        info_box_v.addLayout(col_header)

        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet('color:#e2e8f0')
        info_box_v.addWidget(sep)

        # 스크롤 영역 (동적 행 컨테이너)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet('QScrollArea{background:transparent}')

        self._rows_container = QWidget()
        self._rows_container.setObjectName('rowsContainer')
        self._rows_container.setStyleSheet('#rowsContainer{background:transparent}')
        self._rows_layout = QVBoxLayout(self._rows_container)
        self._rows_layout.setContentsMargins(0, 0, 0, 0)
        self._rows_layout.setSpacing(0)
        self._rows_layout.addStretch()

        scroll.setWidget(self._rows_container)
        info_box_v.addWidget(scroll, 1)

        # 버튼 행
        btn_row = QHBoxLayout()
        add_btn = QPushButton('+ 항목 추가')
        add_btn.setFixedHeight(26)
        add_btn.setStyleSheet(
            'QPushButton{background:#0891b2;color:white;border-radius:4px;font-size:10px;}'
            'QPushButton:hover{background:#0e7490;}'
        )
        add_btn.clicked.connect(lambda: self._add_row())

        reset_btn = QPushButton('기본값 복원')
        reset_btn.setFixedHeight(26)
        reset_btn.setStyleSheet(
            'QPushButton{background:#64748b;color:white;border-radius:4px;font-size:10px;}'
            'QPushButton:hover{background:#475569;}'
        )
        reset_btn.clicked.connect(self._reset_defaults)

        btn_row.addWidget(add_btn)
        btn_row.addWidget(reset_btn)
        btn_row.addStretch()
        info_box_v.addLayout(btn_row)

        lv.addWidget(info_box, 1)

        # 기본 항목 로드
        self._reset_defaults()

        # ── 오른쪽: 파일 목록 ────────────────────────────────────
        right = QWidget()
        rv    = QVBoxLayout(right)
        rv.setContentsMargins(6, 0, 0, 0)
        rv.setSpacing(8)

        file_box = QGroupBox("📁 장비 파일 목록")
        fv = QVBoxLayout(file_box)

        fbtn_row = QHBoxLayout()
        self.btn_add = QPushButton("+ 파일 추가")
        self.btn_add.setFixedHeight(28)
        self.btn_add.clicked.connect(self._add_files)
        self.btn_del = QPushButton("− 선택 제거")
        self.btn_del.setFixedHeight(28)
        self.btn_del.clicked.connect(self._remove_selected)
        self.btn_clr = QPushButton("전체 제거")
        self.btn_clr.setFixedHeight(28)
        self.btn_clr.clicked.connect(self._clear_files)
        fbtn_row.addWidget(self.btn_add)
        fbtn_row.addWidget(self.btn_del)
        fbtn_row.addWidget(self.btn_clr)
        fbtn_row.addStretch()
        fv.addLayout(fbtn_row)

        self.file_list = QListWidget()
        self.file_list.setFont(QFont("Consolas", 9))
        self.file_list.setSelectionMode(QListWidget.ExtendedSelection)
        fv.addWidget(self.file_list)

        hint = QLabel("※ show version / show proc cpu / show proc memory / show dir (flash) / show logging 출력이 포함된 txt 파일")
        hint.setWordWrap(True)
        hint.setStyleSheet("color:#64748b;font-size:10px")
        fv.addWidget(hint)

        rv.addWidget(file_box, 1)

        splitter.addWidget(left)
        splitter.addWidget(right)
        splitter.setSizes([300, 680])
        bv.addWidget(splitter, 1)

        # ── 하단: 진행 바 + 생성 버튼 ────────────────────────────
        self.progress = QProgressBar()
        self.progress.setFixedHeight(5)
        self.progress.setTextVisible(False)
        self.progress.setRange(0, 0)
        self.progress.hide()
        bv.addWidget(self.progress)

        btn_bar = QHBoxLayout()
        self.btn_pdf   = QPushButton("📄 PDF 보고서")
        self.btn_word  = QPushButton("📝 Word 보고서")
        self.btn_excel = QPushButton("📊 Excel 보고서")
        for b in (self.btn_pdf, self.btn_word, self.btn_excel):
            b.setFixedHeight(34)
            b.setFont(QFont("맑은 고딕", 10, QFont.Bold))
        self.btn_pdf.setStyleSheet("background:#2563eb;color:#fff;border-radius:6px")
        self.btn_word.setStyleSheet("background:#16a34a;color:#fff;border-radius:6px")
        self.btn_excel.setStyleSheet("background:#0891b2;color:#fff;border-radius:6px")
        self.btn_pdf.clicked.connect(lambda: self._generate('pdf'))
        self.btn_word.clicked.connect(lambda: self._generate('word'))
        self.btn_excel.clicked.connect(lambda: self._generate('excel'))

        self.lbl_status = QLabel("파일을 추가하고 보고서를 생성하세요")
        self.lbl_status.setStyleSheet("color:#64748b;font-size:11px")

        btn_bar.addWidget(self.btn_pdf)
        btn_bar.addWidget(self.btn_word)
        btn_bar.addWidget(self.btn_excel)
        btn_bar.addStretch()
        btn_bar.addWidget(self.lbl_status)
        bv.addLayout(btn_bar)

    # ── 동적 행 관리 ──────────────────────────────────────────────
    def _add_row(self, label='', value='', placeholder=''):
        row = _InfoRow(label, value, placeholder)
        row.delete_clicked.connect(self._delete_row)
        # 스트레치 앞에 삽입
        idx = self._rows_layout.count() - 1
        self._rows_layout.insertWidget(idx, row)
        self._info_rows.append(row)

    def _delete_row(self, row_widget):
        if row_widget in self._info_rows:
            self._info_rows.remove(row_widget)
        self._rows_layout.removeWidget(row_widget)
        row_widget.deleteLater()

    def _reset_defaults(self):
        # 기존 행 제거
        for row in list(self._info_rows):
            self._rows_layout.removeWidget(row)
            row.deleteLater()
        self._info_rows.clear()
        # 기본 항목 추가
        today = datetime.now().strftime('%Y-%m-%d')
        for label, ph, default in _DEFAULT_FIELDS:
            val = today if label == '점검일' else default
            self._add_row(label, val, ph)

    def _get_info(self):
        items = []
        for row in self._info_rows:
            lbl, val = row.get_pair()
            if lbl:
                items.append({'label': lbl, 'value': val or '-'})
        return {'items': items}

    # ── 파일 추가/제거 ────────────────────────────────────────────
    def _add_files(self):
        paths, _ = QFileDialog.getOpenFileNames(
            self, "show 명령어 파일 선택", "",
            "텍스트 파일 (*.txt *.log *.cfg);;모든 파일 (*)"
        )
        for p in paths:
            existing = [self.file_list.item(i).data(Qt.UserRole)
                        for i in range(self.file_list.count())]
            if p not in existing:
                item = QListWidgetItem(f"  {os.path.basename(p)}")
                item.setData(Qt.UserRole, p)
                item.setToolTip(p)
                self.file_list.addItem(item)
        self.lbl_status.setText(f"{self.file_list.count()}개 파일 등록됨")

    def _remove_selected(self):
        for item in self.file_list.selectedItems():
            self.file_list.takeItem(self.file_list.row(item))
        self.lbl_status.setText(f"{self.file_list.count()}개 파일 등록됨")

    def _clear_files(self):
        self.file_list.clear()
        self.lbl_status.setText("파일을 추가하고 보고서를 생성하세요")

    # ── 보고서 생성 ───────────────────────────────────────────────
    def _generate(self, fmt: str):
        count = self.file_list.count()
        if count == 0:
            QMessageBox.warning(self, "파일 없음", "장비 파일을 먼저 추가하세요.")
            return

        if fmt == 'pdf':
            ext, suffix = "PDF 파일 (*.pdf)", '.pdf'
        elif fmt == 'word':
            ext, suffix = "Word 파일 (*.docx)", '.docx'
        else:
            ext, suffix = "Excel 파일 (*.xlsx)", '.xlsx'

        today   = datetime.now().strftime('%Y%m%d')
        default = f"네트워크점검보고서_{today}"

        save_path, _ = QFileDialog.getSaveFileName(self, "저장 위치", default, ext)
        if not save_path:
            return
        if not save_path.endswith(suffix):
            save_path += suffix

        files = [self.file_list.item(i).data(Qt.UserRole) for i in range(count)]
        info  = self._get_info()

        for b in (self.btn_pdf, self.btn_word, self.btn_excel):
            b.setEnabled(False)
        self.progress.show()

        self._worker = _ReportWorker(fmt, save_path, files, info)
        self._worker.progress.connect(self.lbl_status.setText)
        self._worker.done.connect(self._on_done)
        self._worker.error.connect(self._on_error)
        self._worker.start()

    def _on_done(self, path):
        self.progress.hide()
        for b in (self.btn_pdf, self.btn_word, self.btn_excel):
            b.setEnabled(True)
        self.lbl_status.setText(f"저장 완료: {os.path.basename(path)}")
        QMessageBox.information(self, "완료", f"보고서가 생성되었습니다:\n{path}")

    def _on_error(self, msg):
        self.progress.hide()
        for b in (self.btn_pdf, self.btn_word, self.btn_excel):
            b.setEnabled(True)
        self.lbl_status.setText(f"오류: {msg}")
        QMessageBox.critical(self, "생성 오류", f"보고서 생성 중 오류 발생:\n\n{msg}\n\n"
                             "PDF는 reportlab, Word는 python-docx, Excel은 openpyxl 설치가 필요합니다.")


# ─── 헬퍼: info['items'] → dict ───────────────────────────────────────────
def _info_val(info: dict, label: str, fallback='-') -> str:
    """동적 items 리스트에서 라벨로 값 검색"""
    for item in info.get('items', []):
        if item.get('label') == label:
            return item.get('value', fallback) or fallback
    return fallback


# ═══════════════════════════════════════════════════════════════════════════
# PDF 생성 (reportlab)
# ═══════════════════════════════════════════════════════════════════════════
def _build_pdf(path: str, devices: list, info: dict):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    import platform
    font_name = 'Helvetica'
    if platform.system() == 'Windows':
        candidates = [
            'C:/Windows/Fonts/malgun.ttf',
            'C:/Windows/Fonts/gulim.ttc',
            'C:/Windows/Fonts/batang.ttc',
        ]
        for fp in candidates:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont('KorFont', fp))
                    font_name = 'KorFont'
                except Exception:
                    pass
                break

    W, H = A4
    doc  = SimpleDocTemplate(path, pagesize=A4,
                              leftMargin=20*mm, rightMargin=20*mm,
                              topMargin=20*mm, bottomMargin=20*mm)

    C_DARK   = colors.HexColor('#1e293b')
    C_BLUE   = colors.HexColor('#2563eb')
    C_GRAY   = colors.HexColor('#64748b')
    C_LGRAY  = colors.HexColor('#f8fafc')
    C_BORDER = colors.HexColor('#e2e8f0')
    C_OK     = colors.HexColor('#16a34a')
    C_WARN   = colors.HexColor('#ea580c')
    C_ERR    = colors.HexColor('#dc2626')
    C_THEAD  = colors.HexColor('#1e3a5f')

    def scolor(status):
        return {'정상': C_OK, '주의': C_WARN, '경고': C_ERR}.get(status, C_GRAY)

    def ps(name, base='Normal', **kw):
        kw.setdefault('fontName', font_name)
        return ParagraphStyle(name, parent=getSampleStyleSheet()[base], **kw)

    S_TITLE   = ps('Title2',  fontSize=32, textColor=C_DARK,  alignment=TA_CENTER, spaceAfter=6, leading=40)
    S_SUBTITLE= ps('Sub',     fontSize=13, textColor=C_GRAY,  alignment=TA_CENTER, spaceAfter=4)
    S_H1      = ps('H1',      fontSize=14, textColor=C_DARK,  spaceBefore=10, spaceAfter=6)
    S_H2      = ps('H2',      fontSize=11, textColor=C_BLUE,  spaceBefore=8,  spaceAfter=4)
    S_BODY    = ps('Body',    fontSize=9,  textColor=C_DARK,  leading=14)
    S_SMALL   = ps('Small',   fontSize=8,  textColor=C_GRAY)
    S_CENTER  = ps('Center',  fontSize=9,  textColor=C_DARK,  alignment=TA_CENTER)

    def tbl_style(header=True):
        s = [
            ('FONTNAME',     (0, 0), (-1, -1), font_name),
            ('FONTSIZE',     (0, 0), (-1, -1), 9),
            ('GRID',         (0, 0), (-1, -1), 0.4, C_BORDER),
            ('ROWBACKGROUNDS',(0, 1), (-1, -1), [colors.white, C_LGRAY]),
            ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING',  (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING',   (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
        ]
        if header:
            s += [
                ('BACKGROUND', (0, 0), (-1, 0), C_THEAD),
                ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
                ('FONTNAME',   (0, 0), (-1, 0), font_name),
                ('FONTSIZE',   (0, 0), (-1, 0), 9),
                ('ALIGN',      (0, 0), (-1, 0), 'CENTER'),
            ]
        return TableStyle(s)

    story = []

    # ── 표지 ────────────────────────────────────────────────────
    story.append(Spacer(1, 30*mm))
    story.append(Paragraph("NETWORK INSPECTION REPORT", S_SUBTITLE))
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph("네트워크 점검 보고서", S_TITLE))
    story.append(Spacer(1, 4*mm))
    story.append(HRFlowable(width="100%", thickness=2, color=C_BLUE, spaceAfter=8*mm))

    # 표지에 모든 동적 항목 표시
    cover_data = [[item['label'], item['value']] for item in info.get('items', [])]
    if cover_data:
        ct = Table(cover_data, colWidths=[45*mm, W - 85*mm])
        ct.setStyle(TableStyle([
            ('FONTNAME',     (0, 0), (-1, -1), font_name),
            ('FONTSIZE',     (0, 0), (-1, -1), 11),
            ('FONTSIZE',     (0, 0), (0, -1),  10),
            ('TEXTCOLOR',    (0, 0), (0, -1),  C_GRAY),
            ('TEXTCOLOR',    (1, 0), (1, -1),  C_DARK),
            ('ALIGN',        (0, 0), (0, -1),  'RIGHT'),
            ('VALIGN',       (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',   (0, 0), (-1, -1), 7),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 7),
            ('LINEBELOW',    (0, 0), (-1, -1), 0.3, C_BORDER),
        ]))
        story.append(ct)

    story.append(Spacer(1, 20*mm))
    story.append(Paragraph(f"총 점검 장비 : {len(devices)}대", S_CENTER))
    story.append(PageBreak())

    # ── 문서 정보 ────────────────────────────────────────────────
    story.append(Paragraph("1. 문서 정보", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER, spaceAfter=4*mm))

    info_rows = [['항목', '내용']]
    info_rows += [[item['label'], item['value']] for item in info.get('items', [])]
    info_rows.append(['점검 장비 수', f"{len(devices)} 대"])
    t = Table(info_rows, colWidths=[45*mm, W - 90*mm])
    t.setStyle(tbl_style())
    story.append(t)
    story.append(PageBreak())

    # ── 점검 개요 ────────────────────────────────────────────────
    story.append(Paragraph("2. 점검 개요", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER, spaceAfter=4*mm))

    # No=8 호스트명=42 OS=28 Serial=30 상태=12 비고=나머지(~52mm)
    col_w = [8*mm, 42*mm, 28*mm, 30*mm, 12*mm, 0]
    col_w[-1] = W - 40*mm - sum(col_w[:-1])

    S_BIGO = ps('Bigo', fontSize=8, textColor=C_DARK, leading=11)
    S_BIGO_W = ps('BigoW', fontSize=8, textColor=C_WARN, leading=11)
    S_BIGO_E = ps('BigoE', fontSize=8, textColor=C_ERR, leading=11)

    def bigo_para(d):
        text = '; '.join(d.issues) if d.issues else '-'
        style = {'경고': S_BIGO_E, '주의': S_BIGO_W}.get(d.status, S_BIGO)
        return Paragraph(text, style)

    summary_rows = [['No', '호스트명', 'OS 버전', 'Serial', '상태', '비고']]
    for i, d in enumerate(devices, 1):
        summary_rows.append([
            str(i), d.hostname or '-',
            d.ios_version or '-', d.serial or '-',
            d.status, bigo_para(d),
        ])

    t2 = Table(summary_rows, colWidths=col_w, repeatRows=1)
    base_style = tbl_style()
    base_style.add('VALIGN', (0, 1), (-1, -1), 'TOP')
    for i, d in enumerate(devices, 1):
        base_style.add('TEXTCOLOR', (4, i), (4, i), scolor(d.status))
        base_style.add('FONTNAME',  (4, i), (4, i), font_name)
    t2.setStyle(base_style)
    story.append(t2)
    story.append(Spacer(1, 6*mm))

    legend = Table(
        [['범례', '정상: 이상 없음', '주의: 60~79% 또는 경보 감지', '경고: 80% 이상 또는 다수 경보']],
        colWidths=[14*mm, 38*mm, 62*mm, 55*mm]
    )
    legend.setStyle(TableStyle([
        ('FONTNAME',   (0,0), (-1,-1), font_name),
        ('FONTSIZE',   (0,0), (-1,-1), 8),
        ('TEXTCOLOR',  (1,0), (1,0),   C_OK),
        ('TEXTCOLOR',  (2,0), (2,0),   C_WARN),
        ('TEXTCOLOR',  (3,0), (3,0),   C_ERR),
        ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
        ('BACKGROUND', (0,0), (0,0),   C_LGRAY),
        ('GRID',       (0,0), (-1,-1), 0.3, C_BORDER),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING',(0,0),(-1,-1), 3),
    ]))
    story.append(legend)
    story.append(Spacer(1, 8*mm))

    story.append(Paragraph("■ 점검 결과 요약", S_H2))
    story.append(Spacer(1, 2*mm))

    cnt_ok   = sum(1 for d in devices if d.status == '정상')
    cnt_warn = sum(1 for d in devices if d.status == '주의')
    cnt_err  = sum(1 for d in devices if d.status == '경고')

    stat_tbl = Table(
        [['구분', '장비 수', '비율'],
         ['정상', str(cnt_ok),   f'{cnt_ok/len(devices)*100:.0f}%'   if devices else '-'],
         ['주의', str(cnt_warn), f'{cnt_warn/len(devices)*100:.0f}%' if devices else '-'],
         ['경고', str(cnt_err),  f'{cnt_err/len(devices)*100:.0f}%'  if devices else '-'],
         ['합계', str(len(devices)), '100%']],
        colWidths=[30*mm, 25*mm, 25*mm]
    )
    stat_s = tbl_style()
    stat_s.add('TEXTCOLOR', (0,1), (0,1), C_OK)
    stat_s.add('TEXTCOLOR', (0,2), (0,2), C_WARN)
    stat_s.add('TEXTCOLOR', (0,3), (0,3), C_ERR)
    stat_s.add('BACKGROUND',(0,4), (-1,4), C_LGRAY)
    stat_tbl.setStyle(stat_s)
    story.append(stat_tbl)
    story.append(Spacer(1, 5*mm))

    issue_devs = [d for d in devices if d.status in ('주의', '경고')]
    if issue_devs:
        story.append(Paragraph("■ 조치 필요 장비", S_H2))
        story.append(Spacer(1, 2*mm))
        S_ISS = ps('Iss', fontSize=8, textColor=C_DARK, leading=11)
        iss_sum_rows = [['호스트명', '상태', '주요 이슈']]
        for d in issue_devs:
            iss_text = '; '.join(d.issues) if d.issues else '-'
            iss_sum_rows.append([
                d.hostname or '-',
                d.status,
                Paragraph(iss_text, S_ISS),
            ])
        col_w2 = [45*mm, 12*mm, 0]
        col_w2[-1] = W - 40*mm - sum(col_w2[:-1])
        it = Table(iss_sum_rows, colWidths=col_w2, repeatRows=1)
        iss_s = tbl_style()
        iss_s.add('VALIGN', (0, 1), (-1, -1), 'TOP')
        for i, d in enumerate(issue_devs, 1):
            iss_s.add('TEXTCOLOR', (1, i), (1, i), scolor(d.status))
        it.setStyle(iss_s)
        story.append(it)

    story.append(PageBreak())

    # ── 장비별 상세 ──────────────────────────────────────────────
    story.append(Paragraph("3. 장비별 상세 점검 결과", S_H1))
    story.append(HRFlowable(width="100%", thickness=1, color=C_BORDER, spaceAfter=4*mm))

    for idx, d in enumerate(devices):
        sc = scolor(d.status)
        hdr_tbl = Table(
            [[f"{idx+1}. {d.hostname or d.filename}", f"상태: {d.status}"]],
            colWidths=[W - 80*mm, 30*mm]
        )
        hdr_tbl.setStyle(TableStyle([
            ('FONTNAME',     (0,0), (-1,-1), font_name),
            ('FONTSIZE',     (0,0), (0,0),   12),
            ('FONTSIZE',     (1,0), (1,0),   10),
            ('TEXTCOLOR',    (0,0), (0,0),   C_DARK),
            ('TEXTCOLOR',    (1,0), (1,0),   sc),
            ('ALIGN',        (1,0), (1,0),   'RIGHT'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('BACKGROUND',   (0,0), (-1,-1), C_LGRAY),
            ('TOPPADDING',   (0,0), (-1,-1), 6),
            ('BOTTOMPADDING',(0,0), (-1,-1), 6),
            ('LEFTPADDING',  (0,0), (-1,-1), 8),
            ('LINEBELOW',    (0,0), (-1,-1), 1.5, C_BLUE),
        ]))
        story.append(hdr_tbl)
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ 기본 정보", S_H2))
        bt = Table([
            ['OS 버전',       d.ios_version or '-'],
            ['Serial Number', d.serial or '-'],
            ['업타임',        d.uptime or '-'],
            ['마지막 재시작', d.last_reload_time or '-'],
            ['재시작 원인',   d.reload_reason or '-'],
        ], colWidths=[40*mm, W - 85*mm])
        bt.setStyle(tbl_style(header=False))
        story.append(bt)
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ CPU 현황", S_H2))
        ct2 = Table([['구분', '사용률'],
                     ['최근 5초', d.cpu_5sec or '-'],
                     ['최근 1분', d.cpu_1min or '-'],
                     ['최근 5분', d.cpu_5min or '-']], colWidths=[40*mm, 40*mm])
        ct2.setStyle(tbl_style())
        story.append(ct2)
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ 메모리 현황", S_H2))
        mt = Table([['항목', '값'],
                    ['전체',   d.mem_total_mb],
                    ['여유',   d.mem_free_mb],
                    ['사용률', f'{d.mem_pct:.1f}%' if d.mem_total else '-']],
                   colWidths=[40*mm, 40*mm])
        mt.setStyle(tbl_style())
        story.append(mt)
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ 스토리지 현황", S_H2))
        if d.storages:
            st_rows = [['파일시스템', '전체', '여유', '사용률']]
            for s in d.storages:
                st_rows.append([s.filesystem or '-', s.total_mb, s.free_mb, f'{s.used_pct:.1f}%'])
            stt = Table(st_rows, colWidths=[55*mm, 30*mm, 30*mm, 25*mm])
            stt.setStyle(tbl_style())
            story.append(stt)
        else:
            story.append(Paragraph("정보 없음", S_SMALL))
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ 주요 로그 (ERROR/WARNING 이상)", S_H2))
        if d.notable_logs:
            log_rows = [[Paragraph(l, S_BODY)] for l in d.notable_logs[-20:]]
            lt = Table(log_rows, colWidths=[W - 40*mm])
            lt.setStyle(TableStyle([
                ('FONTNAME',      (0,0), (-1,-1), font_name),
                ('FONTSIZE',      (0,0), (-1,-1), 8),
                ('BACKGROUND',    (0,0), (-1,-1), colors.HexColor('#fff7ed')),
                ('GRID',          (0,0), (-1,-1), 0.3, C_BORDER),
                ('TEXTCOLOR',     (0,0), (-1,-1), colors.HexColor('#7c2d12')),
                ('TOPPADDING',    (0,0), (-1,-1), 3),
                ('BOTTOMPADDING', (0,0), (-1,-1), 3),
                ('LEFTPADDING',   (0,0), (-1,-1), 5),
                ('ROWBACKGROUNDS',(0,0), (-1,-1), [colors.white, colors.HexColor('#fff7ed')]),
            ]))
            story.append(lt)
        else:
            story.append(Paragraph("해당 없음 (주요 경보 미감지)", S_SMALL))
        story.append(Spacer(1, 3*mm))

        story.append(Paragraph("■ 점검 결과 요약", S_H2))
        iss_rows = [[f'• {iss}'] for iss in d.issues] or [['• 이상 없음']]
        ist = Table(iss_rows, colWidths=[W - 40*mm])
        ist.setStyle(TableStyle([
            ('FONTNAME',      (0,0), (-1,-1), font_name),
            ('FONTSIZE',      (0,0), (-1,-1), 9),
            ('TEXTCOLOR',     (0,0), (-1,-1), sc),
            ('BACKGROUND',    (0,0), (-1,-1), C_LGRAY),
            ('GRID',          (0,0), (-1,-1), 0.3, C_BORDER),
            ('TOPPADDING',    (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
            ('LEFTPADDING',   (0,0), (-1,-1), 8),
        ]))
        story.append(ist)

        if idx < len(devices) - 1:
            story.append(PageBreak())

    doc.build(story)


# ═══════════════════════════════════════════════════════════════════════════
# Word 생성 (python-docx)
# ═══════════════════════════════════════════════════════════════════════════
def _build_word(path: str, devices: list, info: dict):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def rgb(hex_str):
        h = hex_str.lstrip('#')
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    C_BLUE  = rgb('#2563eb')
    C_DARK  = rgb('#1e293b')
    C_GRAY  = rgb('#64748b')

    def scolor_rgb(status):
        return {'정상': rgb('#16a34a'), '주의': rgb('#ea580c'), '경고': rgb('#dc2626')}.get(status, C_GRAY)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.name = '맑은 고딕'
        p.runs[0].font.color.rgb = C_DARK
        return p

    def set_cell(cell, text, bold=False, center=False, color=None, bg=None, size=9):
        cell.text = ''
        p   = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  bg.lstrip('#'))
            tcPr.append(shd)

    def add_table(headers, rows, col_widths_cm=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'
        for i, h in enumerate(headers):
            set_cell(t.rows[0].cells[i], h, bold=True, center=True, color=RGBColor(255,255,255), bg='1e3a5f', size=9)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                set_cell(t.rows[r_idx+1].cells[c_idx], str(val), size=9)
        if col_widths_cm:
            for r in t.rows:
                for i, w in enumerate(col_widths_cm):
                    if i < len(r.cells):
                        r.cells[i].width = Cm(w)
        return t

    # 표지
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('네트워크 점검 보고서')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.font.color.rgb = C_DARK

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('NETWORK INSPECTION REPORT').font.color.rgb = C_GRAY

    doc.add_paragraph()
    items = info.get('items', [])
    if items:
        t = doc.add_table(rows=len(items), cols=2)
        t.style = 'Table Grid'
        for i, item in enumerate(items):
            set_cell(t.rows[i].cells[0], item['label'], bold=True, color=C_GRAY, bg='f8fafc', size=10)
            set_cell(t.rows[i].cells[1], item['value'], size=11)
            t.rows[i].cells[0].width = Cm(3.5)
            t.rows[i].cells[1].width = Cm(12)

    doc.add_page_break()

    # 문서 정보
    add_heading('1. 문서 정보', 1)
    info_rows = [[item['label'], item['value']] for item in items]
    info_rows.append(['점검 장비 수', f"{len(devices)} 대"])
    add_table(['항목', '내용'], info_rows, [4, 12])
    doc.add_page_break()

    # 점검 개요
    add_heading('2. 점검 개요', 1)
    summary_rows = []
    for i, d in enumerate(devices, 1):
        summary_rows.append([
            str(i), d.hostname or '-',
            d.ios_version or '-', d.serial or '-',
            d.status, ', '.join(d.issues[:2]) if d.issues else '-',
        ])
    st = add_table(['No', '호스트명', 'OS 버전', 'Serial', '상태', '비고'], summary_rows,
                   [1, 5, 3.5, 3.5, 1.5, 6.5])
    for i, d in enumerate(devices, 1):
        c = st.rows[i].cells[4]
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.color.rgb = scolor_rgb(d.status)

    doc.add_paragraph()
    add_heading('■ 점검 결과 요약', 3)
    cnt_ok   = sum(1 for d in devices if d.status == '정상')
    cnt_warn = sum(1 for d in devices if d.status == '주의')
    cnt_err  = sum(1 for d in devices if d.status == '경고')
    stat_rows = [
        ['정상', str(cnt_ok),   f'{cnt_ok/len(devices)*100:.0f}%'   if devices else '-'],
        ['주의', str(cnt_warn), f'{cnt_warn/len(devices)*100:.0f}%' if devices else '-'],
        ['경고', str(cnt_err),  f'{cnt_err/len(devices)*100:.0f}%'  if devices else '-'],
        ['합계', str(len(devices)), '100%'],
    ]
    st2 = add_table(['구분', '장비 수', '비율'], stat_rows, [3, 2.5, 2.5])
    for i, hex_c in enumerate(['#16a34a', '#ea580c', '#dc2626']):
        c = st2.rows[i+1].cells[0]
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.color.rgb = rgb(hex_c)

    issue_devs = [d for d in devices if d.status in ('주의', '경고')]
    if issue_devs:
        doc.add_paragraph()
        add_heading('■ 조치 필요 장비', 3)
        iss_rows = [[d.hostname or '-', d.status, '; '.join(d.issues[:3])] for d in issue_devs]
        si = add_table(['호스트명', '상태', '주요 이슈'], iss_rows, [5, 1.5, 14.5])
        for i, d in enumerate(issue_devs, 1):
            c = si.rows[i].cells[1]
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.color.rgb = scolor_rgb(d.status)

    doc.add_page_break()

    # 장비별 상세
    add_heading('3. 장비별 상세 점검 결과', 1)
    for idx, d in enumerate(devices):
        p = doc.add_heading(f'{idx+1}. {d.hostname or d.filename}  [상태: {d.status}]', 2)
        for run in p.runs:
            run.font.color.rgb = scolor_rgb(d.status)

        doc.add_heading('■ 기본 정보', 3)
        add_table(['항목', '값'], [
            ['OS 버전',       d.ios_version or '-'],
            ['Serial Number', d.serial or '-'],
            ['업타임',        d.uptime or '-'],
            ['마지막 재시작', d.last_reload_time or '-'],
            ['재시작 원인',   d.reload_reason or '-'],
        ], [4, 12])

        doc.add_heading('■ CPU 현황', 3)
        add_table(['구분', '사용률'], [
            ['최근 5초', d.cpu_5sec or '-'],
            ['최근 1분', d.cpu_1min or '-'],
            ['최근 5분', d.cpu_5min or '-'],
        ], [4, 4])

        doc.add_heading('■ 메모리 현황', 3)
        add_table(['항목', '값'], [
            ['전체',   d.mem_total_mb],
            ['여유',   d.mem_free_mb],
            ['사용률', f'{d.mem_pct:.1f}%' if d.mem_total else '-'],
        ], [4, 4])

        doc.add_heading('■ 스토리지 현황', 3)
        if d.storages:
            add_table(['파일시스템', '전체', '여유', '사용률'],
                      [[s.filesystem or '-', s.total_mb, s.free_mb, f'{s.used_pct:.1f}%']
                       for s in d.storages], [5.5, 3, 3, 2.5])
        else:
            doc.add_paragraph('정보 없음')

        doc.add_heading('■ 주요 로그 (ERROR/WARNING 이상)', 3)
        if d.notable_logs:
            for log in d.notable_logs[-20:]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(log)
                run.font.size = Pt(8)
                run.font.name = 'Consolas'
                run.font.color.rgb = rgb('#7c2d12')
        else:
            doc.add_paragraph('해당 없음 (주요 경보 미감지)')

        doc.add_heading('■ 점검 결과 요약', 3)
        for iss in d.issues:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(iss)
            run.font.color.rgb = scolor_rgb(d.status)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
        if not d.issues:
            doc.add_paragraph('• 이상 없음')

        if idx < len(devices) - 1:
            doc.add_page_break()

    doc.save(path)


# ═══════════════════════════════════════════════════════════════════════════
# Excel 생성 (openpyxl)
# ═══════════════════════════════════════════════════════════════════════════
def _build_excel(path: str, devices: list, info: dict):
    from openpyxl import Workbook
    from openpyxl.styles import (
        Font, PatternFill, Alignment, Border, Side, numbers
    )
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # ── 공통 스타일 헬퍼 ────────────────────────────────────────
    def _fill(hex_color):
        return PatternFill('solid', fgColor=hex_color.lstrip('#'))

    def _font(bold=False, size=10, color='000000', name='맑은 고딕'):
        return Font(bold=bold, size=size, color=color.lstrip('#'), name=name)

    def _border():
        s = Side(style='thin', color='CBD5E1')
        return Border(left=s, right=s, top=s, bottom=s)

    def _align(h='left', v='center', wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    STATUS_COLORS = {'정상': '16a34a', '주의': 'ea580c', '경고': 'dc2626'}

    def _style_row(ws, row_idx, cols, bg=None, bold=False, center=False, color=None):
        for col in range(1, cols + 1):
            cell = ws.cell(row=row_idx, column=col)
            cell.border = _border()
            cell.alignment = _align('center' if center else 'left')
            cell.font = _font(bold=bold, color=color or '000000')
            if bg:
                cell.fill = _fill(bg)

    def _header_row(ws, row_idx, cols):
        _style_row(ws, row_idx, cols, bg='1e3a5f', bold=True, center=True, color='ffffff')

    # ════════════════════════════════════════════
    # 시트 1: 문서 정보
    # ════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = '점검 정보'

    ws1.column_dimensions['A'].width = 20
    ws1.column_dimensions['B'].width = 40

    ws1['A1'] = '네트워크 점검 보고서'
    ws1['A1'].font = _font(bold=True, size=16, color='1e293b')
    ws1['A1'].alignment = _align('left')
    ws1.row_dimensions[1].height = 30
    ws1.merge_cells('A1:B1')

    ws1.append([])

    ws1.append(['항목', '내용'])
    _header_row(ws1, ws1.max_row, 2)

    for item in info.get('items', []):
        ws1.append([item['label'], item['value']])
        r = ws1.max_row
        ws1.cell(r, 1).fill  = _fill('f8fafc')
        ws1.cell(r, 1).font  = _font(bold=True, color='64748b')
        ws1.cell(r, 1).alignment = _align('right')
        ws1.cell(r, 1).border = _border()
        ws1.cell(r, 2).font  = _font()
        ws1.cell(r, 2).alignment = _align()
        ws1.cell(r, 2).border = _border()

    ws1.append(['점검 장비 수', f'{len(devices)} 대'])
    r = ws1.max_row
    ws1.cell(r, 1).fill  = _fill('f8fafc')
    ws1.cell(r, 1).font  = _font(bold=True, color='64748b')
    ws1.cell(r, 1).alignment = _align('right')
    ws1.cell(r, 1).border = _border()
    ws1.cell(r, 2).font  = _font()
    ws1.cell(r, 2).alignment = _align()
    ws1.cell(r, 2).border = _border()

    # ════════════════════════════════════════════
    # 시트 2: 점검 개요
    # ════════════════════════════════════════════
    ws2 = wb.create_sheet('점검 개요')
    col_widths2 = [6, 22, 30, 20, 8, 35]
    for i, w in enumerate(col_widths2, 1):
        ws2.column_dimensions[get_column_letter(i)].width = w

    ws2.append(['No', '호스트명', 'OS 버전', 'Serial', '상태', '비고'])
    _header_row(ws2, 1, 6)
    ws2.row_dimensions[1].height = 20

    for i, d in enumerate(devices, 1):
        ws2.append([
            i, d.hostname or '-',
            d.ios_version or '-', d.serial or '-',
            d.status, ', '.join(d.issues[:2]) if d.issues else '-',
        ])
        r = ws2.max_row
        _style_row(ws2, r, 6, bg='f8fafc' if i % 2 == 0 else None)
        sc = STATUS_COLORS.get(d.status, '6b7280')
        ws2.cell(r, 5).font = _font(bold=True, color=sc)
        ws2.cell(r, 5).alignment = _align('center')

    # ════════════════════════════════════════════
    # 시트 3: 점검 결과 요약
    # ════════════════════════════════════════════
    ws3 = wb.create_sheet('점검 결과 요약')
    ws3.column_dimensions['A'].width = 12
    ws3.column_dimensions['B'].width = 10
    ws3.column_dimensions['C'].width = 10

    ws3.append(['구분', '장비 수', '비율'])
    _header_row(ws3, 1, 3)

    cnt_ok   = sum(1 for d in devices if d.status == '정상')
    cnt_warn = sum(1 for d in devices if d.status == '주의')
    cnt_err  = sum(1 for d in devices if d.status == '경고')
    total    = len(devices) or 1

    for label, cnt, color in [
        ('정상', cnt_ok,   '16a34a'),
        ('주의', cnt_warn, 'ea580c'),
        ('경고', cnt_err,  'dc2626'),
        ('합계', len(devices), '1e293b'),
    ]:
        pct = f'{cnt/total*100:.0f}%'
        ws3.append([label, cnt, pct])
        r = ws3.max_row
        _style_row(ws3, r, 3, center=True)
        ws3.cell(r, 1).font = _font(bold=(label == '합계'), color=color)

    # ════════════════════════════════════════════
    # 시트 4+: 장비별 상세
    # ════════════════════════════════════════════
    for d in devices:
        title = (d.hostname or d.filename or 'unknown')[:25]
        ws = wb.create_sheet(title)
        ws.column_dimensions['A'].width = 22
        ws.column_dimensions['B'].width = 35

        sc = STATUS_COLORS.get(d.status, '6b7280')

        # 장비 헤더
        ws['A1'] = f'{d.hostname or d.filename}  [상태: {d.status}]'
        ws['A1'].font = _font(bold=True, size=13, color=sc)
        ws['A1'].alignment = _align()
        ws.row_dimensions[1].height = 25
        ws.merge_cells('A1:B1')
        ws.append([])

        def _section(title_txt, rows):
            ws.append([title_txt])
            r = ws.max_row
            ws.cell(r, 1).font = _font(bold=True, size=10, color='2563eb')
            ws.cell(r, 1).alignment = _align()
            ws.merge_cells(f'A{r}:B{r}')

            ws.append(['항목', '값'])
            r2 = ws.max_row
            _header_row(ws, r2, 2)

            for k, v in rows:
                ws.append([k, str(v)])
                rr = ws.max_row
                ws.cell(rr, 1).fill  = _fill('f8fafc')
                ws.cell(rr, 1).font  = _font(color='475569')
                ws.cell(rr, 1).alignment = _align('right')
                ws.cell(rr, 1).border = _border()
                ws.cell(rr, 2).font  = _font()
                ws.cell(rr, 2).alignment = _align()
                ws.cell(rr, 2).border = _border()
            ws.append([])

        _section('■ 기본 정보', [
            ('OS 버전',       d.ios_version or '-'),
            ('Serial Number', d.serial or '-'),
            ('업타임',        d.uptime or '-'),
            ('마지막 재시작', d.last_reload_time or '-'),
            ('재시작 원인',   d.reload_reason or '-'),
        ])

        _section('■ CPU 현황', [
            ('최근 5초', d.cpu_5sec or '-'),
            ('최근 1분', d.cpu_1min or '-'),
            ('최근 5분', d.cpu_5min or '-'),
        ])

        _section('■ 메모리 현황', [
            ('전체',   d.mem_total_mb),
            ('여유',   d.mem_free_mb),
            ('사용률', f'{d.mem_pct:.1f}%' if d.mem_total else '-'),
        ])

        # 스토리지
        ws.append(['■ 스토리지 현황'])
        r = ws.max_row
        ws.cell(r, 1).font = _font(bold=True, size=10, color='2563eb')
        ws.merge_cells(f'A{r}:D{r}')
        for col_i, header in enumerate(['파일시스템', '전체', '여유', '사용률'], 1):
            ws.cell(r + 1, col_i).value = header
        r2 = ws.max_row + 1
        ws.append(['파일시스템', '전체', '여유', '사용률'])
        _header_row(ws, ws.max_row, 4)
        for col_i in [3, 4]:
            ws.column_dimensions[get_column_letter(col_i)].width = 14
        if d.storages:
            for s in d.storages:
                ws.append([s.filesystem or '-', s.total_mb, s.free_mb, f'{s.used_pct:.1f}%'])
                rr = ws.max_row
                _style_row(ws, rr, 4)
        else:
            ws.append(['정보 없음', '', '', ''])
        ws.append([])

        # 주요 로그
        ws.append(['■ 주요 로그 (ERROR/WARNING 이상)'])
        r = ws.max_row
        ws.cell(r, 1).font = _font(bold=True, size=10, color='2563eb')
        ws.merge_cells(f'A{r}:B{r}')
        ws.column_dimensions['A'].width = 70
        ws.column_dimensions['B'].width = 5
        if d.notable_logs:
            for log in d.notable_logs[-30:]:
                ws.append([log])
                rr = ws.max_row
                ws.cell(rr, 1).font = Font(size=8, name='Consolas', color='7c2d12')
                ws.cell(rr, 1).fill = _fill('fff7ed')
                ws.cell(rr, 1).alignment = _align(wrap=True)
                ws.cell(rr, 1).border = _border()
                ws.row_dimensions[rr].height = 14
        else:
            ws.append(['해당 없음'])
        ws.append([])

        # 점검 결과 요약
        ws.append(['■ 점검 결과 요약'])
        r = ws.max_row
        ws.cell(r, 1).font = _font(bold=True, size=10, color='2563eb')
        ws.merge_cells(f'A{r}:B{r}')
        for iss in (d.issues or ['이상 없음']):
            ws.append([f'• {iss}'])
            rr = ws.max_row
            ws.cell(rr, 1).font = _font(color=sc)
            ws.cell(rr, 1).fill = _fill('f8fafc')
            ws.cell(rr, 1).border = _border()

    wb.save(path)
